import streamlit as st
from docx import Document
from io import BytesIO
import base64
from zhipuai import ZhipuAI
from concurrent.futures import ThreadPoolExecutor, as_completed

st.title("商业计划书生成器")

company = st.text_input("公司名称")
describe = st.text_area("一段话描述公司（行业，愿景使命）")
product = st.text_area("产品介绍")
team = st.text_area("团队成员及分工")

# Function to generate business plan
def generate_business_plan(company, describe, product, team):

    client = ZhipuAI(api_key="")  # 填写您自己的APIKey

    def get_task_prompt(company, role, role_describe, describe, product, team, bussiness_content, wrok_content,
                        content_title, word_limit=500):
        content = f"""
        #角色
        你是{company}公司的{role}

        #角色描述
        {role_describe}

        #公司背景信息
        {company}
        {describe}

        #公司主打产品
        {product}

        #公司成员
        {team}

        #商业计划书相关部分内容：
        {bussiness_content}

        #任务
        现在你需要进行关于商业计划书{content_title}的撰写工作，具体工作如下：
        {wrok_content}

        #原则
        1.请说得更具体一点。要有创造力和想象力。
        2.请用{word_limit}个字回复指定的任务。不要添加任何其他不存在的内容，永远不要设定虚假时间。

        """
        return content

    def get_task(Messages):
        response = client.chat.completions.create(
            model="glm-4",  # 填写需要调用的模型名称
            messages=Messages,
        )
        res = response.choices[0].message.content
        return res

    # 企业描述
    def CompanyDescription_task():
        role = "创始人"
        role_describe = """作为商业计划书的主要发起者和推动者，创始人负责整体战略规划、愿景和目标的确定，
        并与团队共同制定商业计划书的方向和内容。他们也负责商业模式的设计和核心业务概念的提出。
        """

        content_title = "企业描述"
        wrok_content = """
        - 公司使命和愿景陈述，强调技术创新和解决方案提供。
        - 公司产品介绍。
        - 业务结构，着重介绍技术团队和研发部门。
        """

        bussiness_content = ""
        CompanyDescription_Prompt = get_task_prompt(company, role, role_describe, describe, product, team,
                                                    bussiness_content, wrok_content, content_title, word_limit=500)
        CompanyDescription_Messages = [{"role": "user", "content": CompanyDescription_Prompt}]
        CompanyDescription = get_task(CompanyDescription_Messages)
        return CompanyDescription

    # 市场分析
    def Market_Analysis_task():
        role = "CMO"
        role_describe = """负责收集和分析市场数据、行业趋势以及竞争情报，并提供市场分析和竞争对手分析的内容。负责制定市场推广策略、销售渠道管理以及品牌推广计划，提供市场定位、推广活动等内容。
        """

        content_title = "市场分析"
        wrok_content = """
       - 行业概况和趋势，重点突出技术驱动的行业发展趋势和创新方向。
    - 目标市场的描述，强调技术产品或解决方案的市场需求和定位。
    - 竞争对手分析和定位，特别关注竞争对手的技术实力和创新能力。
        """

        bussiness_content = ""
        Market_Analysis_Prompt = get_task_prompt(company, role, role_describe, describe, product, team,
                                                 bussiness_content,
                                                 wrok_content, content_title, word_limit=500)
        Market_Analysis_Messages = [{"role": "user", "content": Market_Analysis_Prompt}]
        Market_Analysis = get_task(Market_Analysis_Messages)
        return Market_Analysis

    # 组织与管理
    def Organization_Management_task():
        role = "创始人"
        role_describe = ""

        content_title = "组织与管理"
        wrok_content = """
    - 公司组织结构，突出技术团队和研发部门在组织结构中的地位和作用。
    - 高层管理团队的介绍，强调技术领导人员的背景和经验。
    - 工作职责和责任，明确技术团队的职责和技术开发任务。
        """

        bussiness_content = ""
        Organization_Management_Prompt = get_task_prompt(company, role, role_describe, describe, product, team,
                                                         bussiness_content, wrok_content, content_title, word_limit=500)
        Organization_Management_Messages = [{"role": "user", "content": Organization_Management_Prompt}]
        Organization_Management = get_task(Organization_Management_Messages)
        return Organization_Management

    # 技术描述
    def Technical_Description_task():
        role = "创始人"
        role_describe = "负责技术产品或解决方案的研发规划和实施，提供技术描述、研发计划和技术创新方案等内容。"

        content_title = "技术描述"
        wrok_content = """
    - 技术产品或解决方案的详细描述，包括技术架构、功能特点、技术优势等。
    - 技术开发和研究成果，介绍技术团队的研发背景和技术创新。
    - 技术路线图，说明技术产品或解决方案的开发计划和未来发展方向。
        """

        bussiness_content = ""
        Technical_Description_Prompt = get_task_prompt(company, role, role_describe, describe, product, team,
                                                       bussiness_content, wrok_content, content_title, word_limit=500)
        Technical_Description_Messages = [{"role": "user", "content": Technical_Description_Prompt}]
        Technical_Description = get_task(Technical_Description_Messages)
        return Technical_Description

    # 产品描述（Products）
    def Products_Description_task(bussiness_content):
        role = "产品经理"
        role_describe = "负责产品策划和管理，与技术团队合作定义产品功能和特性，提供产品描述和开发计划。"

        content_title = "技术描述"
        wrok_content = """
    - 技术产品或解决方案的详细描述，突出其技术特点和创新之处。
    - 售价策略，根据技术产品或解决方案的市场定位和价值确定定价策略。
    - 研发和生产计划，说明技术产品或解决方案的研发流程、时间表和资源投入。
        """

        bussiness_content = ""
        Products_Description_Prompt = get_task_prompt(company, role, role_describe, describe, product, team,
                                                      bussiness_content, wrok_content, content_title, word_limit=500)
        Products_Description_Messages = [{"role": "user", "content": Products_Description_Prompt}]
        Products_Description = get_task(Products_Description_Messages)
        return Products_Description

    # 市场推广策略（Marketing and Sales Strategy）
    def Marketing_Strategy_task(bussiness_content):
        role = "CMO"
        role_describe = """负责收集和分析市场数据、行业趋势以及竞争情报，并提供市场分析和竞争对手分析的内容。负责制定市场推广策略、销售渠道管理以及品牌推广计划，提供市场定位、推广活动等内容。
        """
        content_title = "市场推广策略"
        wrok_content = """
    - 市场定位和目标客户群，强调技术产品或解决方案的市场定位和目标客户群体。
    - 销售和营销渠道，根据技术产品或解决方案的特点选择适合的销售和营销渠道。
    - 宣传推广活动，结合技术特点和创新优势设计宣传推广活动和营销策略。
        """

        Marketing_Strategy_Prompt = get_task_prompt(company, role, role_describe, describe, product, team,
                                                    bussiness_content, wrok_content, content_title, word_limit=500)
        Marketing_Strategy_Messages = [{"role": "user", "content": Marketing_Strategy_Prompt}]
        Marketing_Strategy = get_task(Marketing_Strategy_Messages)
        return Marketing_Strategy

    # 财务计划（Financial_Projections）
    def Financial_Projections_task(bussiness_content):
        role = "CFO"
        role_describe = """负责财务规划、预测和资金管理，提供财务数据、预测和资金筹集计划等内容。负责进行财务分析，包括收入和支出预测、现金流量分析、盈利和损益预测等。
        """
        content_title = "财务计划"
        wrok_content = """
    - 销售预测，基于技术产品或解决方案的市场需求和定位进行销售额预测。
    - 收入和支出预测，考虑技术研发和产品推广的成本进行收入和支出预测。
    - 现金流量预测，确保技术公司具有足够的现金流来支持技术开发和商业运营。
    - 盈利和损益表，评估技术公司的盈利能力和财务健康状况。
        """
        Financial_Projections_Prompt = get_task_prompt(company, role, role_describe, describe, product, team,
                                                       bussiness_content, wrok_content, content_title, word_limit=500)
        Financial_Projections_Messages = [{"role": "user", "content": Financial_Projections_Prompt}]
        Financial_Projections = get_task(Financial_Projections_Messages)
        return Financial_Projections

    # 融资计划（Funding_Request）
    def Funding_Request_task(bussiness_content):
        role = "CFO兼法律顾问"
        role_describe = ""
        content_title = "融资计划"
        wrok_content = """
    - 资金需求和用途说明，突出技术研发和产品推广的资金需求。
    - 资金筹集计划，说明技术公司的融资计划和资金筹集渠道。
    - 商业合同、知识产权保护、法律风险评估等，提供法律条款、合规建议和法律风险分析。
        """
        Funding_Request_Prompt = get_task_prompt(company, role, role_describe, describe, product, team,
                                                 bussiness_content,
                                                 wrok_content, content_title, word_limit=500)
        Funding_Request_Messages = [{"role": "user", "content": Funding_Request_Prompt}]
        Funding_Request = get_task(Funding_Request_Messages)
        return Funding_Request

    # 执行摘要（Executive_Summary）
    def Executive_Summary_task(bussiness_content):
        role = "创始人"
        role_describe = ""
        content_title = "执行摘要"
        wrok_content = """
        -总结整个计划的概要。
    - 介绍业务概念，包括产品或服务的核心特点、目标市场和技术优势。
    - 提供清晰的市场机会描述，包括市场规模、增长趋势、竞争分析和技术创新。
    - 概述公司目标和策略，特别强调技术创新对公司发展的重要性。
        """
        Executive_Summary_Prompt = get_task_prompt(company, role, role_describe, describe, product, team,
                                                   bussiness_content,
                                                   wrok_content, content_title, word_limit=500)
        Executive_Summary_Messages = [{"role": "user", "content": Executive_Summary_Prompt}]
        Executive_Summary = get_task(Executive_Summary_Messages)
        return Executive_Summary

    with ThreadPoolExecutor(max_workers=4) as executor:
        future_to_task = {
            executor.submit(CompanyDescription_task): "CompanyDescription",
            executor.submit(Market_Analysis_task): "Market_Analysis",
            executor.submit(Organization_Management_task): "Organization_Management",
            executor.submit(Technical_Description_task): "Technical_Description",
        }

        results = {}

        for future in as_completed(future_to_task):
            task_name = future_to_task[future]
            try:
                result = future.result()
                results[task_name] = result
            except Exception as exc:
                print(f"{task_name} generated an exception: {exc}")
            else:
                print(f"{task_name} completed successfully.")

    CompanyDescription = results["CompanyDescription"]
    Market_Analysis = results["Market_Analysis"]
    Organization_Management = results["Organization_Management"]
    Technical_Description = results["Technical_Description"]

    # 产品描述（Products）
    bussiness_content = "#### 企业描述:\n" + CompanyDescription + "\n" + "#### 市场分析:\n" + Market_Analysis + "\n" + "#### 组织与管理:\n" + Organization_Management + "\n" + "#### 技术描述:\n" + Technical_Description
    Products_Description = Products_Description_task(bussiness_content)

    # 市场推广策略（Marketing and Sales Strategy）
    bussiness_content = "#### 企业描述:\n" + CompanyDescription + "\n" + "#### 市场分析:\n" + Market_Analysis + "\n" + "#### 组织与管理:\n" + Organization_Management + "\n" + "#### 技术描述:\n" + Technical_Description + "\n" + "#### 产品描述:\n" + Products_Description
    Marketing_Strategy = Marketing_Strategy_task(bussiness_content)

    # 财务计划（Financial Projections）
    bussiness_content = "#### 企业描述:\n" + CompanyDescription + "\n" + "#### 市场分析:\n" + Market_Analysis + "\n" + "#### 组织与管理:\n" + Organization_Management + "\n" + "#### 技术描述:\n" + Technical_Description + "\n" + "#### 产品描述:\n" + Products_Description + "\n" + "#### 市场推广策略:\n" + Marketing_Strategy
    Financial_Projections = Financial_Projections_task(bussiness_content)

    # 融资计划
    bussiness_content = "#### 市场推广策略:\n" + Marketing_Strategy + "\n" + "#### 财务计划:\n" + Financial_Projections
    Funding_Request = Funding_Request_task(bussiness_content)

    # 执行摘要
    bussiness_content = "#### 企业描述:\n" + CompanyDescription + "\n" + "#### 市场分析:\n" + Market_Analysis + "\n" + "#### 组织与管理:\n" + Organization_Management + "\n" + "#### 技术描述:\n" + Technical_Description + "\n" + "#### 产品描述:\n" + Products_Description + "\n" + "#### 市场推广策略:\n" + Marketing_Strategy + "\n" + "#### 财务计划:\n" + Financial_Projections + "\n" + "#### 融资计划:\n" + Funding_Request
    Executive_Summary = Executive_Summary_task(bussiness_content)

    doc = Document()

    # Add headings and content to the document
    doc.add_heading('一、执行摘要:', level=2)
    doc.add_paragraph(Executive_Summary)

    doc.add_heading('二、企业描述:', level=2)
    doc.add_paragraph(CompanyDescription)

    doc.add_heading('三、市场分析:', level=2)
    doc.add_paragraph(Market_Analysis)

    doc.add_heading('四、组织与管理:', level=2)
    doc.add_paragraph(Organization_Management)

    doc.add_heading('五、技术描述:', level=2)
    doc.add_paragraph(Technical_Description)

    doc.add_heading('六、产品描述:', level=2)
    doc.add_paragraph(Products_Description)

    doc.add_heading('七、市场推广策略:', level=2)
    doc.add_paragraph(Marketing_Strategy)

    doc.add_heading('八、财务计划:', level=2)
    doc.add_paragraph(Financial_Projections)

    doc.add_heading('九、融资计划:', level=2)
    doc.add_paragraph(Funding_Request)

    # Save the document
    doc.save('business_plan.docx')
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Function to create a download link for the document
    def get_download_link(docx_buffer, filename='business_plan.docx'):
        docx_bytes = docx_buffer.getvalue()
        b64 = base64.b64encode(docx_bytes).decode()
        return f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">点击此链接下载商业计划书</a>'


    return Executive_Summary,CompanyDescription,Market_Analysis,Organization_Management,Technical_Description,Products_Description,Marketing_Strategy,Financial_Projections,Funding_Request,buffer, get_download_link(buffer)



# Button to generate business plan
if st.button("生成商业计划书"):
    Executive_Summary,CompanyDescription,Market_Analysis,Organization_Management,Technical_Description,Products_Description,Marketing_Strategy,Financial_Projections,Funding_Request,buffer, download_link = generate_business_plan(company, describe, product, team)
    st.success("商业计划书生成成功！")

    # Display the download link
    st.markdown(download_link, unsafe_allow_html=True)

    st.markdown("## 商业计划书内容")
    st.markdown("#### 一、执行摘要:\n"+Executive_Summary+"\n"+"#### 二、企业描述:\n"+CompanyDescription+"\n"+"#### 三、市场分析:\n"+Market_Analysis+"\n"+"#### 四、组织与管理:\n"+Organization_Management+"\n"+"#### 五、技术描述:\n"+Technical_Description+"\n"+"#### 六、产品描述:\n"+Products_Description+"\n"+"#### 七、市场推广策略:\n"+Marketing_Strategy+"\n"+"#### 八、财务计划:\n"+Financial_Projections+"\n"+"#### 九、融资计划:\n"+Funding_Request)
